"""
Text-extraction module.

Responsible ONLY for turning uploaded file bytes into clean plain text.
Knows nothing about Gemini, Mongo, or HTTP — kept isolated so it can be
unit-tested and swapped out independently.
"""
import io
import os
import re
from dataclasses import dataclass

from pypdf import PdfReader
import docx  # python-docx

from app.config import settings


class ExtractionError(Exception):
    """Raised for any problem turning a file into usable text."""


@dataclass
class ExtractedDocument:
    text: str
    char_count: int
    page_or_paragraph_count: int
    source_type: str  # "pdf" | "docx"


def _clean_text(raw: str) -> str:
    """Normalize whitespace and drop control characters that sometimes
    leak out of PDF extraction."""
    if not raw:
        return ""
    # Remove null bytes / non-printable control chars (keep newlines/tabs)
    raw = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", raw)
    # Collapse 3+ blank lines into 2, collapse runs of spaces
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


def extract_text_from_pdf(file_bytes: bytes) -> ExtractedDocument:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not open PDF file: {e}") from e

    if reader.is_encrypted:
        try:
            reader.decrypt("")  # try empty password, common for "locked" exports
        except Exception:
            raise ExtractionError(
                "This PDF is password-protected and cannot be read."
            )

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            # Skip unreadable pages rather than failing the whole document
            continue

    text = _clean_text("\n".join(pages_text))
    return ExtractedDocument(
        text=text,
        char_count=len(text),
        page_or_paragraph_count=len(reader.pages),
        source_type="pdf",
    )


def extract_text_from_docx(file_bytes: bytes) -> ExtractedDocument:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not open DOCX file: {e}") from e

    parts = [p.text for p in document.paragraphs]

    # Tables often hold contact info / skills grids — don't skip them.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    text = _clean_text("\n".join(parts))
    return ExtractedDocument(
        text=text,
        char_count=len(text),
        page_or_paragraph_count=len(document.paragraphs),
        source_type="docx",
    )


def extract_text(filename: str, file_bytes: bytes) -> ExtractedDocument:
    """Dispatch based on file extension, then run validation checks."""
    ext = os.path.splitext(filename.lower())[1]

    if ext not in settings.ALLOWED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file type '{ext}'. Only PDF and DOCX are accepted."
        )

    if ext == ".pdf":
        doc = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        doc = extract_text_from_docx(file_bytes)
    else:  # pragma: no cover - guarded above
        raise ExtractionError(f"Unsupported file type '{ext}'.")

    _validate_extracted_text(doc)
    return doc


def _validate_extracted_text(doc: ExtractedDocument) -> None:
    """Reject empty or garbled extraction before it ever reaches Gemini —
    saves an API call and gives the user an actionable error instead of a
    confusing low score."""
    if doc.char_count < settings.MIN_EXTRACTED_TEXT_CHARS:
        raise ExtractionError(
            "Very little readable text could be extracted from this file. "
            "It may be a scanned image, empty, or corrupted. Try exporting "
            "a text-based PDF or DOCX instead."
        )

    # Garbled-text heuristic: proportion of "normal" characters
    # (letters, digits, common punctuation, whitespace) should dominate.
    normal_chars = re.findall(r"[A-Za-z0-9\s.,;:@()/#&+%'\"-]", doc.text)
    ratio = len(normal_chars) / max(doc.char_count, 1)
    if ratio < 0.7:
        raise ExtractionError(
            "The extracted text looks garbled (likely a scanned/image-based "
            "document without embedded text). Please upload a text-based "
            "PDF or DOCX, or an OCR'd version."
        )
